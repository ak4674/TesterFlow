import json
from docx import Document
from docx.shared import RGBColor, Pt
from pathlib import Path


def export_txt(transcript: dict, path: str) -> str:
    with open(path, "w", encoding="utf-8") as f:
        f.write(transcript["text"])
    return path


def export_json(transcript: dict, path: str) -> str:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(transcript, f, ensure_ascii=False, indent=2)
    return path


def export_docx(transcript: dict, path: str) -> str:
    doc = Document()
    doc.add_heading("Transcript", 0)

    meta = doc.add_paragraph()
    meta.add_run(f"Language detected: {transcript.get('language', 'unknown').upper()}  |  "
                 f"Duration: {transcript.get('duration', 0):.1f}s")
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    for seg in transcript.get("segments", []):
        ts = doc.add_paragraph()
        ts.add_run(f"[{_fmt_time(seg['start'])} → {_fmt_time(seg['end'])}]  ")
        ts.runs[-1].font.size = Pt(8)
        ts.runs[-1].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        for word_obj in seg.get("words", []):
            run = ts.add_run(word_obj["word"])
            run.font.size = Pt(11)
            if word_obj["script"] == "devanagari":
                run.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)  # blue for Hindi
            else:
                run.font.color.rgb = RGBColor(0x10, 0x10, 0x10)  # near-black for English

    doc.save(path)
    return path


def _fmt_time(seconds: float) -> str:
    m, s = divmod(int(seconds), 60)
    return f"{m:02d}:{s:02d}"
